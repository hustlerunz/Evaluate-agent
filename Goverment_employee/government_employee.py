from tkinter import N
from tkinter.messagebox import NO
import docx
import pandas as pd
from docx import Document
import openpyxl
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH 
from docx.enum.text import WD_UNDERLINE
import configparser
import sys 
from docx.shared import RGBColor

config = configparser.ConfigParser()
config.read('of.ini')
source_file_data = config['file']['source_data']
master_file = config['file']['form_master']
row_run = int(config['file']['ro_start'])
ch = 0

book = openpyxl.load_workbook(source_file_data)
sheet = book.active
row_count = sheet.max_row

sara = ['่','้','๊','๋','ิ','ี','ื','ุ','ู','์','ั']
sara2 = ['โ','ไ','เ','ใ','.']
#print(sara)

print("excel_sum_row = ",row_count)  
print("========run========")  
for x in range(row_count-1):
    chek_len_count = 0 #check char
    chek_len_count2 = 0 #check char
    var2_chk_half = 0
    var1_chk_half = 0
    #print(x)
    chk_name = sheet.cell(row=row_run+x, column=3).value
    if chk_name == None:
        book.close()
        sys.exit()

    doc = Document(master_file)
    name = sheet.cell(row=row_run+x, column=3).value
    sname = sheet.cell(row=row_run+x, column=4).value
    position = sheet.cell(row=row_run+x, column=5).value
    type = sheet.cell(row=row_run+x, column=6).value
    own = sheet.cell(row=row_run+x, column=7).value

    st_point = sheet.cell(row=row_run+x,column=8).value
    st_point2 = sheet.cell(row=row_run+x,column=9).value

    l1_los = sheet.cell(row=row_run+x,column=10).value
    l1_los_e = sheet.cell(row=row_run+x,column=11).value

    l1_late = sheet.cell(row=row_run+x,column=12).value
    l1_late_e = sheet.cell(row=row_run+x,column=13).value
    l1_bus = sheet.cell(row=row_run+x,column=14).value
    l1_bus_e = sheet.cell(row=row_run+x,column=15).value
    l1_sick = sheet.cell(row=row_run+x,column=16).value
    l1_sick_e = sheet.cell(row=row_run+x,column=17).value
    l1_ordan = sheet.cell(row=row_run+x,column=18).value
    l1_ordan_e = sheet.cell(row=row_run+x,column=19).value
    l1_born = sheet.cell(row=row_run+x,column=20).value	    
    l1_born_e = sheet.cell(row=row_run+x,column=21).value	    
    l1_other = sheet.cell(row=row_run+x,column=22).value	    
    l1_other_e = sheet.cell(row=row_run+x,column=23).value	    
    l1_sum = sheet.cell(row=row_run+x,column=24).value
    l1_sum_e = sheet.cell(row=row_run+x,column=25).value

    l2_los = sheet.cell(row=row_run+x,column=26).value
    l2_los_e = sheet.cell(row=row_run+x,column=27).value
    l2_late = sheet.cell(row=row_run+x,column=28).value
    l2_late_e = sheet.cell(row=row_run+x,column=29).value
    l2_bus = sheet.cell(row=row_run+x,column=30).value
    l2_bus_e = sheet.cell(row=row_run+x,column=31).value
    l2_sick = sheet.cell(row=row_run+x,column=32).value
    l2_sick_e = sheet.cell(row=row_run+x,column=33).value
    l2_ordan = sheet.cell(row=row_run+x,column=34).value
    l2_ordan_e = sheet.cell(row=row_run+x,column=35).value
    l2_born = sheet.cell(row=row_run+x,column=36).value
    l2_born_e = sheet.cell(row=row_run+x,column=37).value
    l2_other = sheet.cell(row=row_run+x,column=38).value
    l2_other_e = sheet.cell(row=row_run+x,column=39).value
    l2_sum = sheet.cell(row=row_run+x,column=40).value
    l2_sum_e = sheet.cell(row=row_run+x,column=41).value

    var1 = '   '+name+'  '+sname+'   ' 
    var2 = '   '+position+'   '
    var4 = '   '+type+'   '
    var3 = ' '
    var33 = '.'
    var5 = '  '+own+'  '

    def chk_sp(var):
        var1_chk_half = 0
        chek_len_count = 0
        for len_var in var:
            for chk in sara:
                if chk == len_var:
                    chek_len_count = chek_len_count+1
            for cha in sara2:
                if cha == len_var:
                    var1_chk_half = var1_chk_half+1
        sum_sp =48-(len(var)-chek_len_count)-(var1_chk_half/2)
        return sum_sp                    
   
    print("var1_lenth = ",chk_sp(var1))
    print("var2_lenth = ",chk_sp(var2))
    print("var3_lenth = ",chk_sp(var4))
    print("var4_lenth = ",chk_sp(var5))
    #print("var2_lenth = ",len(chk_len2))                        
    print("var1_upper_under = ",chek_len_count)            
    print("var2_upper_under = ",chek_len_count2)
    print("var1_half_c = ",var1_chk_half)            
    print("var2_half_c = ",var2_chk_half)                 
    
    print("var1_sum_postion = ",((len(var1)-chek_len_count)-chek_len_count2/2))

    newdata = []
    newdata2 = []
    var_string = " "
    
   
    def insert_space(sp_num,name_sname):
        for x in range(int(sp_num)):
            name_sname += ' '
        return(name_sname)    
    
    print(insert_space(chk_sp(var1),var1))
    
    var_name = doc.tables[0].cell(0,0).paragraphs[1].text = insert_space(chk_sp(var1),var1)
    var_name1 = doc.tables[0].cell(0,0).paragraphs[1].add_run('.')
    var_name1.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
 
    var_postion = doc.tables[0].cell(0,1).paragraphs[1].text = insert_space(chk_sp(var2),var2)
    var_name1 = doc.tables[0].cell(0,1).paragraphs[1].add_run('.')
    var_name1.font.color.rgb = RGBColor(0xff, 0xff, 0xff)

    var_postion = doc.tables[0].cell(1,0).paragraphs[1].text = insert_space(chk_sp(var4),var4)
    var_name1 = doc.tables[0].cell(1,0).paragraphs[1].add_run('.')
    var_name1.font.color.rgb = RGBColor(0xff, 0xff, 0xff)

    var_postion = doc.tables[0].cell(1,1).paragraphs[1].text = insert_space(chk_sp(var5),var5)
    var_name1 = doc.tables[0].cell(1,1).paragraphs[1].add_run('.')
    var_name1.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
    
    if int(st_point) >= 90: 
        v12 = doc.tables[1].cell(2,1).paragraphs[0].text = str(st_point)
        doc.tables[1].cell(2,1).paragraphs[0].alignment = 1
        newdata2.append(doc.tables[1].cell(2,1).paragraphs[0])
    elif st_point >= 80 and st_point <=89:
        print(st_point)
        v12 = doc.tables[1].cell(2,2).paragraphs[0].text = str(st_point)
        doc.tables[1].cell(2,2).paragraphs[0].alignment = 1
        newdata2.append(doc.tables[1].cell(2,2).paragraphs[0])
    elif int(st_point) >= 70 and int(st_point) <=79:
        v12 = doc.tables[1].cell(2,3).paragraphs[0].text = str(st_point)
        doc.tables[1].cell(2,3).paragraphs[0].alignment = 1
        newdata2.append(doc.tables[1].cell(2,3).paragraphs[0])
    elif int(st_point) >= 60 and int(st_point) <=69:
        v12 = doc.tables[1].cell(2,4).paragraphs[0].text = str(st_point)
        doc.tables[1].cell(2,4).paragraphs[0].alignment = 1
        newdata2.append(doc.tables[1].cell(2,4).paragraphs[0])
    else:
        v12 = doc.tables[1].cell(2,5).paragraphs[0].text = str(st_point)
        doc.tables[1].cell(2,5).paragraphs[0].alignment = 1
        newdata2.append(doc.tables[1].cell(2,5).paragraphs[0])
    #round2  
      
    if st_point2 != None:
        if int(st_point2) >= 90: 
            v2_12 = doc.tables[1].cell(3,1).paragraphs[0].text = str(st_point2)
            doc.tables[1].cell(3,1).paragraphs[0].alignment = 1
            newdata2.append(doc.tables[1].cell(3,1).paragraphs[0])
        
        elif st_point2 >= 80 and st_point2 <=89:
            v2_12 = doc.tables[1].cell(3,2).paragraphs[0].text = str(st_point2)
            doc.tables[1].cell(3,2).paragraphs[0].alignment = 1
            newdata2.append(doc.tables[1].cell(3,2).paragraphs[0])
        elif int(st_point2) >= 70 and int(st_point2) <=79:
            v2_12 = doc.tables[1].cell(3,3).paragraphs[0].text = str(st_point2)
            doc.tables[1].cell(3,3).paragraphs[0].alignment = 1
            newdata.append(doc.tables[0].cell(3,3).paragraphs[0])
        elif int(st_point2) >= 60 and int(st_point2) <=69:
            v2_12 = doc.tables[1].cell(3,4).paragraphs[0].text = str(st_point2)
            doc.tables[1].cell(3,4).paragraphs[0].alignment = 1
            newdata.append(doc.tables[1].cell(3,4).paragraphs[0])
        else:
            v2_12 = doc.tables[1].cell(3,5).paragraphs[0].text = str(st_point2)
            doc.tables[1].cell(3,5).paragraphs[0].alignment = 1
            newdata.append(doc.tables[1].cell(3,5).paragraphs[0]) 

    v13 = doc.tables[2].cell(2,2).paragraphs[0].text = str(l1_los)    
    v14 = doc.tables[2].cell(2,3).paragraphs[0].text = str(l1_los_e)    
    v15 = doc.tables[2].cell(3,2).paragraphs[0].text = str(l1_late)
    v16 = doc.tables[2].cell(3,3).paragraphs[0].text = str(l1_late_e)
    v17 = doc.tables[2].cell(4,2).paragraphs[0].text = str(l1_bus)
    v18 = doc.tables[2].cell(4,3).paragraphs[0].text = str(l1_bus_e)
    v19 = doc.tables[2].cell(5,2).paragraphs[0].text = str(l1_sick)
    v20 = doc.tables[2].cell(5,3).paragraphs[0].text = str(l1_sick_e)
    v21 = doc.tables[2].cell(6,2).paragraphs[0].text = str(l1_ordan)
    v22 = doc.tables[2].cell(6,3).paragraphs[0].text = str(l1_ordan_e)
    v23 = doc.tables[2].cell(7,2).paragraphs[0].text = str(l1_born)
    v24 = doc.tables[2].cell(7,3).paragraphs[0].text = str(l1_born_e)
    v25 = doc.tables[2].cell(8,2).paragraphs[0].text = str(l1_other)
    v26 = doc.tables[2].cell(8,3).paragraphs[0].text = str(l1_other_e)
    v27 = doc.tables[2].cell(9,2).paragraphs[0].text = str(l1_sum)
    v28 = doc.tables[2].cell(9,3).paragraphs[0].text = str(l1_sum_e)

    v29 = doc.tables[2].cell(2,4).paragraphs[0].text = str(l2_los) 
    v30 = doc.tables[2].cell(2,5).paragraphs[0].text = str(l2_los_e)
    v31 = doc.tables[2].cell(3,4).paragraphs[0].text = str(l2_late) 
    v32 = doc.tables[2].cell(3,5).paragraphs[0].text = str(l2_late_e) 
    v33 = doc.tables[2].cell(4,4).paragraphs[0].text = str(l2_bus)
    v34 = doc.tables[2].cell(4,5).paragraphs[0].text = str(l2_bus_e)
    v35 = doc.tables[2].cell(5,4).paragraphs[0].text = str(l2_sick)
    v36 = doc.tables[2].cell(5,5).paragraphs[0].text = str(l2_sick_e)
    v37 = doc.tables[2].cell(6,4).paragraphs[0].text = str(l2_ordan)
    v38 = doc.tables[2].cell(6,5).paragraphs[0].text = str(l2_ordan_e)
    v39 = doc.tables[2].cell(7,4).paragraphs[0].text = str(l2_born)
    v40 = doc.tables[2].cell(7,5).paragraphs[0].text = str(l2_born_e)
    v41 = doc.tables[2].cell(8,4).paragraphs[0].text = str(l2_other)
    v42 = doc.tables[2].cell(8,5).paragraphs[0].text = str(l2_other_e)
    v43 = doc.tables[2].cell(9,4).paragraphs[0].text = str(l2_sum)
    v44 = doc.tables[2].cell(9,5).paragraphs[0].text = str(l2_sum_e)
    
    for x in range(2,10):
        for y in range(2,4):
            newdata2.append(doc.tables[2].cell(x,y).paragraphs[0])
        for y2 in range(4,6):
            newdata2.append(doc.tables[2].cell(x,y2).paragraphs[0])
    newdata.append(doc.tables[0].cell(0,0).paragraphs[1])
    newdata.append(doc.tables[0].cell(0,1).paragraphs[1])
    newdata.append(doc.tables[0].cell(1,0).paragraphs[1])
    newdata.append(doc.tables[0].cell(1,1).paragraphs[1])
    
    for data in newdata:
        data.runs[0].underline = True
        data.runs[0].underline = WD_UNDERLINE.DOTTED
        data.runs[0].font.name = 'TH SarabunIT๙'
        data.runs[0].font.size = Pt(16)
    
    for data2 in newdata2:
        data2.runs[0].font.name = 'TH SarabunIT๙'
        data2.runs[0].font.size = Pt(16)

    doc.save(name+'.docx')
    print("========End========")


